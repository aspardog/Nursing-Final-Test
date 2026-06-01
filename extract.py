#!/usr/bin/env python3
"""
extract.py - Extracción mecánica de material de estudio para enfermería
Sin LLM, sin visión, solo código Python determinístico.
"""

import os
from pathlib import Path
import fitz  # pymupdf
from docx import Document
import numpy as np

# Configuración
UPLOADS_DIR = Path("./uploads")
MATERIAL_DIR = Path("./material")
DPI = 200
BLANK_THRESHOLD = 0.98  # 98% píxeles blancos = página en blanco

# Mapeo de archivos a carpetas
ARCHIVOS = {
    "semiologia": {
        "archivo": "3. SEMIOLOGIA PSIQUIATRICA.pdf",
        "carpeta": "01_semiologia",
        "tipo": "pdf_texto",
        "extraer_jpg": False
    },
    "salas": {
        "archivo": "MENTORIA SALAS DE CIRUGÍA.pdf",
        "carpeta": "02_salas_cirugia",
        "tipo": "pdf_mixto",
        "extraer_jpg": True
    },
    "urgencias": {
        "archivo": "MENTORIA URGENCIAS.pdf",
        "carpeta": "03_urgencias",
        "tipo": "pdf_mixto",
        "extraer_jpg": True
    },
    "arritmias": {
        "archivo": "Arritmias cardíacas_250707_121615.pdf",
        "carpeta": "04_arritmias",
        "tipo": "pdf_imagen",
        "extraer_jpg": True
    },
    "mezclas": {
        "archivo": "EJERCICIOS DE INTRODUCCIÓN DE MEZCLAS Y DILUCIONES.docx",
        "carpeta": "05_mezclas",
        "tipo": "docx",
        "extraer_jpg": False
    }
}

# Almacenar resultados para el resumen
resultados = []
errores = []


def crear_estructura_carpetas():
    """Crea la estructura de carpetas necesaria."""
    print("\n[SETUP] Creando estructura de carpetas...")

    for key, config in ARCHIVOS.items():
        carpeta = MATERIAL_DIR / config["carpeta"]
        carpeta.mkdir(parents=True, exist_ok=True)

        if config["extraer_jpg"]:
            jpg_dir = carpeta / "paginas_jpg"
            jpg_dir.mkdir(exist_ok=True)

    print(f"[OK] Estructura creada en {MATERIAL_DIR}/")


def diagnosticar_pdf(ruta: Path) -> tuple:
    """
    Diagnostica un PDF: número de páginas y si tiene texto extraíble.
    Retorna: (num_paginas, tiene_texto, texto_por_pagina)
    """
    doc = fitz.open(ruta)
    num_paginas = len(doc)

    texto_total = ""
    texto_por_pagina = []

    for page in doc:
        texto = page.get_text().strip()
        texto_por_pagina.append(texto)
        texto_total += texto

    tiene_texto = len(texto_total.strip()) > 100  # Más de 100 chars = tiene texto
    doc.close()

    return num_paginas, tiene_texto, texto_por_pagina


def es_pagina_blanca(pixmap) -> bool:
    """
    Determina si una página es mayormente blanca.
    """
    # Convertir pixmap a array numpy
    img_data = pixmap.samples
    img_array = np.frombuffer(img_data, dtype=np.uint8)

    # Reshape según el número de componentes (RGB o RGBA)
    n_components = pixmap.n
    img_array = img_array.reshape(-1, n_components)

    # Calcular píxeles "casi blancos" (RGB > 250)
    if n_components >= 3:
        blancos = np.all(img_array[:, :3] > 250, axis=1)
    else:
        blancos = img_array[:, 0] > 250

    porcentaje_blanco = np.mean(blancos)
    return porcentaje_blanco > BLANK_THRESHOLD


def extraer_texto_pdf(ruta: Path, carpeta_destino: Path, guardar_individual: bool = True) -> dict:
    """
    Extrae texto de un PDF.
    Retorna: dict con info de extracción
    """
    doc = fitz.open(ruta)
    num_paginas = len(doc)

    texto_completo = ""
    archivos_generados = []

    for i, page in enumerate(doc):
        texto = page.get_text()
        texto_completo += f"\n{'='*60}\n"
        texto_completo += f"PÁGINA {i+1}\n"
        texto_completo += f"{'='*60}\n"
        texto_completo += texto

        if guardar_individual:
            archivo_pagina = carpeta_destino / f"pagina_{i+1:02d}.txt"
            with open(archivo_pagina, "w", encoding="utf-8") as f:
                f.write(texto)
            archivos_generados.append(str(archivo_pagina.relative_to(MATERIAL_DIR)))

    # Guardar texto completo
    archivo_completo = carpeta_destino / "texto_completo.txt"
    with open(archivo_completo, "w", encoding="utf-8") as f:
        f.write(texto_completo)
    archivos_generados.insert(0, str(archivo_completo.relative_to(MATERIAL_DIR)))

    doc.close()

    return {
        "paginas": num_paginas,
        "caracteres": len(texto_completo),
        "archivos": archivos_generados
    }


def extraer_texto_pdf_mixto(ruta: Path, carpeta_destino: Path) -> dict:
    """
    Extrae texto de un PDF mixto (Canva) - guarda texto.txt + páginas individuales.
    """
    doc = fitz.open(ruta)
    num_paginas = len(doc)

    texto_completo = ""
    archivos_generados = []

    for i, page in enumerate(doc):
        texto = page.get_text()
        texto_completo += f"\n{'='*60}\n"
        texto_completo += f"PÁGINA {i+1}\n"
        texto_completo += f"{'='*60}\n"
        texto_completo += texto

        archivo_pagina = carpeta_destino / f"pagina_{i+1:02d}.txt"
        with open(archivo_pagina, "w", encoding="utf-8") as f:
            f.write(texto)
        archivos_generados.append(str(archivo_pagina.relative_to(MATERIAL_DIR)))

    # Guardar texto completo como texto.txt
    archivo_texto = carpeta_destino / "texto.txt"
    with open(archivo_texto, "w", encoding="utf-8") as f:
        f.write(texto_completo)
    archivos_generados.insert(0, str(archivo_texto.relative_to(MATERIAL_DIR)))

    doc.close()

    return {
        "paginas": num_paginas,
        "caracteres": len(texto_completo),
        "archivos": archivos_generados
    }


def pdf_a_jpg(ruta: Path, carpeta_destino: Path, detectar_blancos: bool = False) -> dict:
    """
    Convierte páginas de PDF a JPG a 200 DPI.
    Si detectar_blancos=True, omite páginas en blanco.
    """
    doc = fitz.open(ruta)
    num_paginas = len(doc)

    jpg_dir = carpeta_destino / "paginas_jpg"
    archivos_generados = []
    paginas_blancas = []
    paginas_con_contenido = 0

    # Matrix para 200 DPI (72 es el DPI base de PDF)
    matrix = fitz.Matrix(DPI / 72, DPI / 72)

    for i, page in enumerate(doc):
        pixmap = page.get_pixmap(matrix=matrix)

        if detectar_blancos and es_pagina_blanca(pixmap):
            paginas_blancas.append(i + 1)
            print(f"    [SKIP] Página {i+1} detectada como en blanco")
            continue

        paginas_con_contenido += 1
        archivo_jpg = jpg_dir / f"pagina_{i+1:02d}.jpg"
        pixmap.save(str(archivo_jpg))
        archivos_generados.append(str(archivo_jpg.relative_to(MATERIAL_DIR)))

    doc.close()

    return {
        "paginas_totales": num_paginas,
        "paginas_con_contenido": paginas_con_contenido,
        "paginas_blancas": paginas_blancas,
        "archivos": archivos_generados
    }


def extraer_docx(ruta: Path, carpeta_destino: Path) -> dict:
    """
    Extrae texto de un archivo DOCX preservando saltos de línea.
    """
    doc = Document(ruta)

    lineas = []
    for para in doc.paragraphs:
        lineas.append(para.text)

    # También extraer texto de tablas si las hay
    for table in doc.tables:
        for row in table.rows:
            fila_texto = []
            for cell in row.cells:
                fila_texto.append(cell.text.strip())
            lineas.append(" | ".join(fila_texto))

    texto_completo = "\n".join(lineas)

    archivo_destino = carpeta_destino / "ejercicios.txt"
    with open(archivo_destino, "w", encoding="utf-8") as f:
        f.write(texto_completo)

    return {
        "elementos": len(doc.paragraphs) + sum(len(t.rows) for t in doc.tables),
        "caracteres": len(texto_completo),
        "archivos": [str(archivo_destino.relative_to(MATERIAL_DIR))]
    }


def generar_resumen():
    """Genera material/resumen.md con tabla de resultados."""
    print("\n[RESUMEN] Generando resumen.md...")

    contenido = "# Resumen de Extracción\n\n"
    contenido += "| Archivo Original | Tipo Detectado | Páginas/Elementos | Archivos Generados | Caracteres |\n"
    contenido += "|-----------------|----------------|-------------------|-------------------|------------|\n"

    for r in resultados:
        archivos_str = ", ".join(r["archivos"][:3])
        if len(r["archivos"]) > 3:
            archivos_str += f" (+{len(r['archivos'])-3} más)"

        contenido += f"| {r['nombre']} | {r['tipo']} | {r['paginas']} | {archivos_str} | {r['caracteres']} |\n"

    if errores:
        contenido += "\n## Errores\n\n"
        for err in errores:
            contenido += f"- **{err['archivo']}**: {err['error']}\n"

    archivo_resumen = MATERIAL_DIR / "resumen.md"
    with open(archivo_resumen, "w", encoding="utf-8") as f:
        f.write(contenido)

    print(f"[OK] Resumen guardado en {archivo_resumen}")


def procesar_semiologia():
    """Procesa el PDF de Semiología Psiquiátrica."""
    config = ARCHIVOS["semiologia"]
    ruta = UPLOADS_DIR / config["archivo"]
    carpeta = MATERIAL_DIR / config["carpeta"]

    print(f"\n[PROC] Procesando: {config['archivo']}")

    # Diagnóstico
    num_pag, tiene_texto, _ = diagnosticar_pdf(ruta)
    texto_status = "SÍ" if tiene_texto else "NO"
    print(f"[DIAG] {config['archivo']}: {num_pag} páginas, texto extraíble: {texto_status}")

    # Extracción
    resultado = extraer_texto_pdf(ruta, carpeta, guardar_individual=True)

    resultados.append({
        "nombre": config["archivo"],
        "tipo": "PDF-texto",
        "paginas": resultado["paginas"],
        "archivos": resultado["archivos"],
        "caracteres": resultado["caracteres"]
    })

    print(f"[OK] Extraídas {resultado['paginas']} páginas, {resultado['caracteres']} caracteres")
    return True


def procesar_salas():
    """Procesa el PDF de Salas de Cirugía (Canva mixto)."""
    config = ARCHIVOS["salas"]
    ruta = UPLOADS_DIR / config["archivo"]
    carpeta = MATERIAL_DIR / config["carpeta"]

    print(f"\n[PROC] Procesando: {config['archivo']}")

    # Diagnóstico
    num_pag, tiene_texto, _ = diagnosticar_pdf(ruta)
    texto_status = "SÍ" if tiene_texto else "NO"
    print(f"[DIAG] {config['archivo']}: {num_pag} páginas, texto extraíble: {texto_status}")

    # Extracción de texto
    resultado_texto = extraer_texto_pdf_mixto(ruta, carpeta)

    # Conversión a JPG
    resultado_jpg = pdf_a_jpg(ruta, carpeta, detectar_blancos=False)

    archivos_todos = resultado_texto["archivos"] + resultado_jpg["archivos"]

    resultados.append({
        "nombre": config["archivo"],
        "tipo": "PDF-mixto",
        "paginas": resultado_texto["paginas"],
        "archivos": archivos_todos,
        "caracteres": resultado_texto["caracteres"]
    })

    print(f"[OK] Extraídas {resultado_texto['paginas']} páginas texto + {len(resultado_jpg['archivos'])} JPGs")
    return True


def procesar_urgencias():
    """Procesa el PDF de Urgencias (Canva mixto)."""
    config = ARCHIVOS["urgencias"]
    ruta = UPLOADS_DIR / config["archivo"]
    carpeta = MATERIAL_DIR / config["carpeta"]

    print(f"\n[PROC] Procesando: {config['archivo']}")

    # Diagnóstico
    num_pag, tiene_texto, _ = diagnosticar_pdf(ruta)
    texto_status = "SÍ" if tiene_texto else "NO"
    print(f"[DIAG] {config['archivo']}: {num_pag} páginas, texto extraíble: {texto_status}")

    # Extracción de texto
    resultado_texto = extraer_texto_pdf_mixto(ruta, carpeta)

    # Conversión a JPG
    resultado_jpg = pdf_a_jpg(ruta, carpeta, detectar_blancos=False)

    archivos_todos = resultado_texto["archivos"] + resultado_jpg["archivos"]

    resultados.append({
        "nombre": config["archivo"],
        "tipo": "PDF-mixto",
        "paginas": resultado_texto["paginas"],
        "archivos": archivos_todos,
        "caracteres": resultado_texto["caracteres"]
    })

    print(f"[OK] Extraídas {resultado_texto['paginas']} páginas texto + {len(resultado_jpg['archivos'])} JPGs")
    return True


def procesar_arritmias():
    """Procesa el PDF de Arritmias (imágenes escaneadas)."""
    config = ARCHIVOS["arritmias"]
    ruta = UPLOADS_DIR / config["archivo"]
    carpeta = MATERIAL_DIR / config["carpeta"]

    print(f"\n[PROC] Procesando: {config['archivo']}")

    # Diagnóstico
    num_pag, tiene_texto, _ = diagnosticar_pdf(ruta)
    texto_status = "SÍ" if tiene_texto else "NO"
    print(f"[DIAG] {config['archivo']}: {num_pag} páginas, texto extraíble: {texto_status}")

    # Solo conversión a JPG (detectando blancos)
    print("    Detectando páginas en blanco...")
    resultado_jpg = pdf_a_jpg(ruta, carpeta, detectar_blancos=True)

    paginas_info = f"{resultado_jpg['paginas_con_contenido']}/{resultado_jpg['paginas_totales']}"
    if resultado_jpg["paginas_blancas"]:
        paginas_info += f" (blancas: {resultado_jpg['paginas_blancas']})"

    resultados.append({
        "nombre": config["archivo"],
        "tipo": "PDF-imagen",
        "paginas": paginas_info,
        "archivos": resultado_jpg["archivos"],
        "caracteres": 0
    })

    print(f"[OK] Convertidas {resultado_jpg['paginas_con_contenido']} páginas a JPG")
    if resultado_jpg["paginas_blancas"]:
        print(f"    Páginas en blanco omitidas: {resultado_jpg['paginas_blancas']}")
    return True


def procesar_mezclas():
    """Procesa el DOCX de Mezclas y Diluciones."""
    config = ARCHIVOS["mezclas"]
    ruta = UPLOADS_DIR / config["archivo"]
    carpeta = MATERIAL_DIR / config["carpeta"]

    print(f"\n[PROC] Procesando: {config['archivo']}")
    print(f"[DIAG] {config['archivo']}: DOCX")

    resultado = extraer_docx(ruta, carpeta)

    resultados.append({
        "nombre": config["archivo"],
        "tipo": "DOCX",
        "paginas": resultado["elementos"],
        "archivos": resultado["archivos"],
        "caracteres": resultado["caracteres"]
    })

    print(f"[OK] Extraídos {resultado['elementos']} elementos, {resultado['caracteres']} caracteres")
    return True


def main():
    print("=" * 60)
    print("EXTRACCIÓN MECÁNICA DE MATERIAL DE ENFERMERÍA")
    print("=" * 60)

    # Crear estructura
    crear_estructura_carpetas()

    # Procesar cada archivo con manejo de errores
    procesadores = [
        ("Semiología", procesar_semiologia),
        ("Salas de Cirugía", procesar_salas),
        ("Urgencias", procesar_urgencias),
        ("Arritmias", procesar_arritmias),
        ("Mezclas", procesar_mezclas),
    ]

    exitos = []

    for nombre, procesador in procesadores:
        try:
            exito = procesador()
            exitos.append((nombre, exito, None))
        except Exception as e:
            print(f"[ERROR] {nombre}: {str(e)}")
            errores.append({"archivo": nombre, "error": str(e)})
            exitos.append((nombre, False, str(e)))

    # Generar resumen
    generar_resumen()

    # Reporte final
    print("\n" + "=" * 60)
    print("REPORTE FINAL")
    print("=" * 60)

    for nombre, exito, error in exitos:
        if exito:
            print(f"✅ {nombre}")
        else:
            print(f"❌ {nombre}: {error}")

    print("\n[FIN] Extracción completada.")


if __name__ == "__main__":
    main()
